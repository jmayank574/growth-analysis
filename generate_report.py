import pandas as pd
import numpy as np
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_PATH   = r"C:\Users\Mayank Joshi\Downloads\Marketing_Channel_Project\data"
OUTPUT_PATH = r"C:\Users\Mayank Joshi\Downloads\Marketing_Channel_Project"

# ── LOAD ALL DATA ─────────────────────────────────────────────────────────────
master       = pd.read_csv(os.path.join(DATA_PATH, 'master_table.csv'))
attribution  = pd.read_csv(os.path.join(DATA_PATH, 'channel_attribution', 'attribution_comparison.csv'))
ltv_channel  = pd.read_csv(os.path.join(DATA_PATH, 'cac_ltv', 'ltv_by_channel.csv'))
ltv_segment  = pd.read_csv(os.path.join(DATA_PATH, 'cac_ltv', 'ltv_by_segment.csv'))
activated    = pd.read_csv(os.path.join(DATA_PATH, 'cac_ltv', 'activated_sellers_ltv.csv'))
rfm          = pd.read_csv(os.path.join(DATA_PATH, 'segmentation', 'rfm_scores.csv'))
segments     = pd.read_csv(os.path.join(DATA_PATH, 'segmentation', 'sellers_segmented_final.csv'))
lookalike    = pd.read_csv(os.path.join(DATA_PATH, 'segmentation', 'high_value_cluster_final.csv'))
ch_funnel    = pd.read_csv(os.path.join(DATA_PATH, 'funnel_analysis', 'channel_funnel.csv'))
cohort       = pd.read_csv(os.path.join(DATA_PATH, 'funnel_analysis', 'cohort_monthly.csv'))

master['first_contact_date'] = pd.to_datetime(master['first_contact_date'], errors='coerce')
master['won_date']           = pd.to_datetime(master['won_date'], errors='coerce')

# ── PRE-COMPUTE KEY METRICS ───────────────────────────────────────────────────
total_leads      = len(master)
total_converted  = int(master['converted'].sum())
total_activated  = int((master['total_orders'] > 0).sum())
total_revenue    = master['total_revenue'].sum()
cvr              = round(total_converted / total_leads * 100, 2)
activation_rate  = round(total_activated / total_converted * 100, 1)
never_activated  = total_converted - total_activated
avg_ltv          = round(activated['total_revenue'].mean(), 2)
median_ltv       = round(activated['total_revenue'].median(), 2)

act_sorted = activated.sort_values('total_revenue', ascending=False).copy()
act_sorted['cumulative_revenue'] = act_sorted['total_revenue'].cumsum()
act_sorted['cumulative_pct']     = act_sorted['cumulative_revenue'] / act_sorted['total_revenue'].sum() * 100
act_sorted['seller_pct']         = np.arange(1, len(act_sorted)+1) / len(act_sorted) * 100
pareto_threshold = act_sorted[act_sorted['cumulative_pct'] >= 80].iloc[0]
pareto_pct       = round(pareto_threshold['seller_pct'], 1)

best_cvr_channel  = ch_funnel.sort_values('cvr', ascending=False).iloc[0]
best_ltv_channel  = ltv_channel.sort_values('avg_ltv', ascending=False).iloc[0]
best_ltv_segment  = ltv_segment.sort_values('avg_ltv', ascending=False).iloc[0]
top_lookalike_ch  = lookalike['origin'].value_counts().index[0]
top_lookalike_seg = lookalike['business_segment'].value_counts().index[0]

rfm_summary = (
    rfm.groupby('rfm_segment')
    .agg(sellers=('seller_id','count'), avg_revenue=('monetary','mean'), avg_orders=('frequency','mean'))
    .reset_index().sort_values('avg_revenue', ascending=False)
)

cluster_summary = (
    segments.groupby('cluster_label')
    .agg(sellers=('seller_id','count'), avg_revenue=('monetary','mean'), total_revenue=('monetary','sum'))
    .reset_index()
)
cluster_summary['revenue_share'] = (
    cluster_summary['total_revenue'] / cluster_summary['total_revenue'].sum() * 100
).round(1)
cluster_summary = cluster_summary.sort_values('avg_revenue', ascending=False)

# ── HELPER FUNCTIONS ──────────────────────────────────────────────────────────
GREEN  = RGBColor(0x0F, 0x6E, 0x56)
BLACK  = RGBColor(0x1A, 0x1A, 0x18)
GRAY   = RGBColor(0x6C, 0x75, 0x7D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LTGREEN= RGBColor(0xE1, 0xF5, 0xEE)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = GREEN if level == 1 else BLACK
        run.font.bold      = True
        run.font.size      = Pt(16) if level == 1 else Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size  = Pt(10.5)
        run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(6)
    return p

def add_insight(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(8)
    border = OxmlElement('w:pBdr')
    left   = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '0F6E56')
    border.append(left)
    p._p.get_or_add_pPr().append(border)
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.italic    = True
    run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style         = 'Table Grid'
    table.alignment     = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, '0F6E56')
        run = cell.paragraphs[0].runs[0]
        run.font.bold      = True
        run.font.color.rgb = WHITE
        run.font.size      = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9.5)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'F0FDF9')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table

def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and bold_part in text:
        parts = text.split(bold_part, 1)
        run1 = p.add_run(parts[0])
        run1.font.size = Pt(10.5)
        run2 = p.add_run(bold_part)
        run2.font.bold = True
        run2.font.size = Pt(10.5)
        if len(parts) > 1:
            run3 = p.add_run(parts[1])
            run3.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

# ── BUILD DOCUMENT ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Growth Analytics Report')
run.font.size      = Pt(28)
run.font.bold      = True
run.font.color.rgb = GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Olist Seller Acquisition — End-to-End Analysis')
run2.font.size      = Pt(14)
run2.font.color.rgb = GRAY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
run3.font.size      = Pt(11)
run3.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

# ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────────
add_heading(doc, '1. Executive Summary', level=1)
add_body(doc,
    f"This report presents a comprehensive growth analytics study of Olist's seller "
    f"acquisition funnel using {total_leads:,} marketing qualified leads from June 2017 "
    f"to September 2018. The analysis covers multi-touch attribution modeling, full funnel "
    f"performance, unit economics (CAC/LTV), and audience segmentation."
)

doc.add_paragraph()
add_heading(doc, 'Key findings', level=2)

add_bullet(doc,
    f"Overall conversion rate of {cvr}% — only {total_converted:,} of {total_leads:,} leads became sellers.",
    f"{cvr}%")
add_bullet(doc,
    f"Critical activation gap: {never_activated} sellers ({round(never_activated/total_converted*100,1)}%) "
    f"signed deals but never placed a single order.",
    f"{never_activated} sellers")
add_bullet(doc,
    f"Top {pareto_pct}% of active sellers drive 80% of total platform revenue (${total_revenue:,.0f}).",
    f"Top {pareto_pct}%")
add_bullet(doc,
    f"Paid search and organic search together account for 62% of high-value sellers "
    f"despite email and social having higher raw lead volumes.",
    "62% of high-value sellers")
add_bullet(doc,
    f"Email channel is undervalued by 5x under last-touch attribution vs linear model — "
    f"budget decisions based on last-touch alone would incorrectly cut email spend.",
    "5x")
add_bullet(doc,
    f"Conversion rate improved from under 5% in 2017 to consistently above 12% from "
    f"February 2018 — a clear inflection point in sales process maturity.",
    "12%")

doc.add_page_break()

# ── SECTION 2: FUNNEL ─────────────────────────────────────────────────────────
add_heading(doc, '2. Funnel Performance', level=1)
add_body(doc,
    "The seller acquisition funnel spans five stages from initial lead generation "
    "through to repeat selling activity on the platform."
)

funnel_rows = [
    ('Leads generated',            f"{total_leads:,}",       '100.0%',  '—'),
    ('Leads with contact date',     f"{master['first_contact_date'].notna().sum():,}",
                                                              '38.2%',   f"{total_leads - int(master['first_contact_date'].notna().sum()):,} lost"),
    ('Closed deals',               f"{total_converted:,}",   '10.5%',   f"{int(master['first_contact_date'].notna().sum()) - total_converted:,} lost"),
    ('Made first order',           f"{total_activated:,}",   '4.8%',    f"{total_converted - total_activated:,} lost"),
    ('Repeat seller (3+ orders)',  f"{int((master['total_orders']>=3).sum()):,}",
                                                              '3.2%',    f"{total_activated - int((master['total_orders']>=3).sum()):,} lost"),
]
add_table(doc,
    ['Stage', 'Count', '% of Total', 'Drop-off'],
    funnel_rows,
    col_widths=[2.8, 1.0, 1.0, 1.2]
)

add_insight(doc,
    f"The largest absolute drop-off occurs at the top — 4,942 leads have no contact date "
    f"recorded, indicating a CRM tracking gap rather than a sales failure. "
    f"The most critical business problem is the activation gap: "
    f"{never_activated} sellers ({round(never_activated/total_converted*100,1)}%) "
    f"signed deals but never generated revenue."
)

add_heading(doc, 'Cohort trend', level=2)
add_body(doc,
    "Monthly cohort analysis reveals a clear inflection point in early 2018. "
    "Pre-2018 cohorts converted at under 5%. From February 2018 onwards, "
    "the floor conversion rate never dropped below 11.7%, peaking at 18.2% in December 2018."
)

doc.add_page_break()

# ── SECTION 3: ATTRIBUTION ────────────────────────────────────────────────────
add_heading(doc, '3. Channel Attribution Analysis', level=1)
add_body(doc,
    "Three attribution models were built to assess true channel value: "
    "last-touch (100% credit to origin), linear (credit split across landing page touchpoints), "
    "and Markov chain (removal effect — how many conversions would be lost without each channel)."
)

attr_rows = []
for _, row in attribution.sort_values('last_touch_rev', ascending=False).iterrows():
    attr_rows.append((
        row['origin'],
        f"${row['last_touch_rev']:,.0f}",
        f"${row['linear_attributed_revenue']:,.0f}",
        f"${row['markov_attributed_revenue']:,.0f}"
    ))

add_table(doc,
    ['Channel', 'Last touch', 'Linear', 'Markov chain'],
    attr_rows,
    col_widths=[1.8, 1.3, 1.3, 1.3]
)

add_insight(doc,
    "Email is undervalued by 5x under last touch vs linear attribution ($8,485 vs $43,468). "
    "Direct traffic is undervalued by 3x ($21,904 vs $68,306). "
    "Budget decisions based on last-touch alone would incorrectly deprioritise "
    "these channels. The Markov chain model confirms organic search as the most "
    "indispensable channel — its removal effect is the highest of any channel."
)

add_heading(doc, 'Channel CVR and efficiency', level=2)
cvr_rows = []
for _, row in ch_funnel.sort_values('cvr', ascending=False).iterrows():
    cvr_rows.append((
        row['origin'],
        f"{int(row['total_leads']):,}",
        f"{int(row['converted'])}",
        f"{row['cvr']:.1f}%",
        f"{int(row['activated'])}",
        f"${row['revenue_per_lead']:.2f}"
    ))

add_table(doc,
    ['Channel', 'Leads', 'Converted', 'CVR%', 'Activated', 'Rev / Lead'],
    cvr_rows,
    col_widths=[1.6, 0.8, 0.9, 0.7, 0.9, 1.0]
)

doc.add_page_break()

# ── SECTION 4: CAC / LTV ──────────────────────────────────────────────────────
add_heading(doc, '4. CAC / LTV — Unit Economics', level=1)
add_body(doc,
    f"Of the {total_converted:,} converted sellers, only {total_activated:,} ({activation_rate}%) "
    f"placed at least one order. All revenue and LTV calculations are based on this "
    f"activated subset. Total platform revenue generated: ${total_revenue:,.2f}."
)

metrics_rows = [
    ('Total platform revenue',    f"${total_revenue:,.2f}"),
    ('Active sellers',            f"{total_activated:,}"),
    ('Avg LTV per active seller', f"${avg_ltv:,.2f}"),
    ('Median LTV per seller',     f"${median_ltv:,.2f}"),
    ('Activation rate',           f"{activation_rate}%"),
    ('Never activated sellers',   f"{never_activated:,}"),
    ('Leads per active seller',   f"{round(total_leads/total_activated, 1)}"),
]
add_table(doc, ['Metric', 'Value'], metrics_rows, col_widths=[3.0, 2.0])

add_heading(doc, 'LTV by acquisition channel', level=2)
ltv_ch_rows = []
for _, row in ltv_channel.sort_values('avg_ltv', ascending=False).iterrows():
    ltv_ch_rows.append((
        row['origin'],
        f"{int(row['activated_sellers'])}",
        f"${row['avg_ltv']:,.0f}",
        f"${row['median_ltv']:,.0f}",
        f"{row['leads_per_activation']:.0f}",
        f"{row['ltv_to_cac_index']:.2f}"
    ))
add_table(doc,
    ['Channel', 'Active sellers', 'Avg LTV', 'Median LTV', 'Leads / activation', 'LTV/CAC index'],
    ltv_ch_rows,
    col_widths=[1.5, 1.2, 1.0, 1.1, 1.4, 1.2]
)

add_heading(doc, 'LTV by business segment (top 10)', level=2)
ltv_seg_rows = []
for _, row in ltv_segment.head(10).iterrows():
    ltv_seg_rows.append((
        row['business_segment'],
        f"{int(row['sellers'])}",
        f"${row['avg_ltv']:,.0f}",
        f"${row['median_ltv']:,.0f}",
        f"{row['avg_orders']:.1f}"
    ))
add_table(doc,
    ['Segment', 'Sellers', 'Avg LTV', 'Median LTV', 'Avg orders'],
    ltv_seg_rows,
    col_widths=[2.2, 0.9, 1.1, 1.1, 1.0]
)

add_insight(doc,
    f"Pareto analysis shows the top {pareto_pct}% of sellers drive 80% of total revenue. "
    f"This extreme concentration means retaining and acquiring high-LTV sellers "
    f"has a disproportionate impact on platform economics. "
    f"Health_beauty is the most scalable high-LTV segment with {ltv_segment[ltv_segment['business_segment']=='health_beauty']['sellers'].values[0] if 'health_beauty' in ltv_segment['business_segment'].values else 'multiple'} "
    f"active sellers and strong average LTV."
)

doc.add_page_break()

# ── SECTION 5: SEGMENTATION ───────────────────────────────────────────────────
add_heading(doc, '5. Audience Segmentation', level=1)
add_body(doc,
    "Two segmentation approaches were applied to the 380 activated sellers: "
    "RFM scoring (rule-based) and K-means clustering (data-driven). "
    "Both methods converge on the same conclusion — a small group of high-value "
    "sellers drives the majority of platform revenue."
)

add_heading(doc, 'RFM segments', level=2)
rfm_rows = []
for _, row in rfm_summary.iterrows():
    rfm_rows.append((
        row['rfm_segment'],
        f"{int(row['sellers'])}",
        f"${row['avg_revenue']:,.0f}",
        f"{row['avg_orders']:.1f}"
    ))
add_table(doc,
    ['Segment', 'Sellers', 'Avg revenue', 'Avg orders'],
    rfm_rows,
    col_widths=[1.8, 1.0, 1.3, 1.1]
)

add_heading(doc, 'K-means clusters', level=2)
cluster_rows = []
for _, row in cluster_summary.iterrows():
    cluster_rows.append((
        row['cluster_label'],
        f"{int(row['sellers'])}",
        f"${row['avg_revenue']:,.0f}",
        f"{row['revenue_share']:.1f}%"
    ))
add_table(doc,
    ['Cluster', 'Sellers', 'Avg revenue', 'Revenue share'],
    cluster_rows,
    col_widths=[1.8, 1.0, 1.3, 1.2]
)

add_heading(doc, 'High value seller — lookalike profile', level=2)
add_body(doc,
    f"The high-value cluster contains {len(lookalike)} sellers averaging "
    f"${lookalike['monetary'].mean():,.0f} in revenue and {lookalike['frequency'].mean():.1f} orders. "
    f"This profile serves as the targeting specification for future acquisition campaigns."
)

lookalike_rows = [
    ('Cluster size',          f"{len(lookalike)} sellers"),
    ('Avg revenue',           f"${lookalike['monetary'].mean():,.0f}"),
    ('Avg orders',            f"{lookalike['frequency'].mean():.1f}"),
    ('Avg review score',      f"{lookalike['avg_review_score'].mean():.2f} / 5.0"),
    ('Top channel',           f"{top_lookalike_ch} (32.4%)"),
    ('Second channel',        f"organic_search (29.4%)"),
    ('Top segment',           f"{top_lookalike_seg} (20.6%)"),
    ('Business type',         "reseller (94.1%)"),
    ('Lead type',             "online_big (44.1%) + online_medium (35.3%)"),
]
add_table(doc, ['Attribute', 'Value'], lookalike_rows, col_widths=[2.5, 3.5])

doc.add_page_break()

# ── SECTION 6: RECOMMENDATIONS ────────────────────────────────────────────────
add_heading(doc, '6. Recommendations', level=1)

recs = [
    (
        "Redirect email and referral traffic to top-performing landing pages",
        "65.3% of email leads land on low-converting pages (CVR < 3.4%). "
        "Redirecting this traffic to the top 3 landing pages (CVR 19-21%) "
        "could generate ~100 additional conversions from existing lead volume "
        "at zero additional acquisition cost."
    ),
    (
        "Fix the activation gap before increasing acquisition spend",
        f"{never_activated} sellers ({round(never_activated/total_converted*100,1)}%) "
        "signed deals but never placed an order. Improving onboarding for newly "
        "converted sellers would increase revenue without spending a single additional "
        "dollar on marketing. Even a 20% improvement in activation rate would add "
        f"~{round(never_activated * 0.2):,} active sellers to the platform."
    ),
    (
        "Shift attribution model from last-touch to Markov chain for budget decisions",
        "Last-touch attribution undervalues email by 5x and direct traffic by 3x. "
        "Teams making budget decisions based on last-touch data are likely cutting "
        "spend on channels that contribute significantly to assisted conversions."
    ),
    (
        "Target paid_search and organic_search for high-LTV seller acquisition",
        "62% of high-value sellers came through paid_search or organic_search. "
        "Social brings 3x the lead volume of paid_search but produces far fewer "
        "high-LTV sellers. Reallocating 20% of social budget to paid search "
        "would likely improve overall seller quality."
    ),
    (
        "Build a lookalike audience from the high-value cluster profile",
        f"The ideal acquisition target is: paid_search or organic_search channel, "
        f"household_utilities or health_beauty segment, reseller business type, "
        f"online_big or online_medium lead type. This profile should be used as a "
        f"seed audience for Meta and Google lookalike targeting."
    ),
]

for i, (title_text, body_text) in enumerate(recs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {title_text}")
    run.font.bold      = True
    run.font.size      = Pt(11)
    run.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(10)
    add_body(doc, body_text)
    doc.add_paragraph()

# ── FOOTER ────────────────────────────────────────────────────────────────────
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(
    f"Growth Analytics Report  ·  Olist E-Commerce Dataset  ·  "
    f"Generated {datetime.now().strftime('%d %B %Y')}"
)
run.font.size      = Pt(9)
run.font.color.rgb = GRAY

# ── SAVE ──────────────────────────────────────────────────────────────────────
filename = f"Growth_Analytics_Report_{datetime.now().strftime('%d%b%Y')}.docx"
filepath = os.path.join(OUTPUT_PATH, filename)
doc.save(filepath)

print(f"Report generated successfully")
print(f"File: {filename}")
print(f"Path: {filepath}")
print(f"\nSections:")
print(f"  1. Executive summary — 6 key findings")
print(f"  2. Funnel performance — stage table + cohort insight")
print(f"  3. Channel attribution — 3 model comparison + CVR table")
print(f"  4. CAC / LTV — unit economics + segment table")
print(f"  5. Audience segmentation — RFM + K-means + lookalike")
print(f"  6. Recommendations — 5 data-driven actions")